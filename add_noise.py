import random
from docx import Document

input_file = "./data/影响大豆豆粕和豆油价格变动因素.docx"
output_file = "./data/noisy_data.docx"

# noise functions
def add_html_tags(text):
    html_tags = ["<p>", "<div>", "<span>", "<a href='#'>", "<strong>", "</p>", "</div>", "</span>", "</a>", "</strong>"]
    return random.choice(html_tags) + text + random.choice(html_tags)

def add_non_ascii_chars(text):
    non_ascii_chars = ['©', '™', '✓', '☺', '☹', '☀', '☁', '❄', '✈']
    insert_position = random.randint(0, len(text))
    return text[:insert_position] + random.choice(non_ascii_chars) + text[insert_position:]

def add_extra_symbols(text):
    extra_symbols = ['!!', '??', '...', '---', '///']
    return text + " " + random.choice(extra_symbols)

doc = Document(input_file)
new_doc = Document()

for paragraph in doc.paragraphs:
    text = paragraph.text
    if text.strip():
        noisy_text = text
        if random.random() < 0.3:
            noisy_text = add_html_tags(noisy_text)
        if random.random() < 0.3:
            noisy_text = add_non_ascii_chars(noisy_text)
        if random.random() < 0.3:
            noisy_text = add_extra_symbols(noisy_text)

        # duplicate some paragraphs
        new_doc.add_paragraph(noisy_text)
        if random.random() < 0.2:
            new_doc.add_paragraph(noisy_text)

new_doc.save(output_file)
print(f"Data with noise is saved as '{output_file}'.")
