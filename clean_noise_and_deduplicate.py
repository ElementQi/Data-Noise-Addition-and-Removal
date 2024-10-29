import re
from docx import Document

input_file = "./data/noisy_data.docx"
output_file = "./data/cleaned_data.docx"

def remove_html_tags(text):
    return re.sub(r'<[^>]+>', '', text)

def remove_unwanted_chars(text):
    """remove non-Chinese, non-ASCII, just keep the common symbols and space"""
    return re.sub(r'[^\x00-\x7F\u4e00-\u9fa5，。！？【】（）“”‘’；：、\s]', '', text)

def remove_extra_symbols(text):
    return re.sub(r'[!?./-]{2,}', '', text)

def clean_text(text):
    text = remove_html_tags(text)
    text = remove_unwanted_chars(text)
    text = remove_extra_symbols(text)
    return text.strip()

doc = Document(input_file)
new_doc = Document()
seen_paragraphs = set()

for paragraph in doc.paragraphs:
    text = paragraph.text.strip()

    cleaned_text = clean_text(text)

    if cleaned_text and cleaned_text not in seen_paragraphs:
        seen_paragraphs.add(cleaned_text)
        new_doc.add_paragraph(cleaned_text)

new_doc.save(output_file)
print(f"cleaned data saved as '{output_file}'.")